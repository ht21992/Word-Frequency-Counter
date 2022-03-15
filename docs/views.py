from django.shortcuts import render
from django.views.generic import TemplateView
from django.http import HttpResponse, JsonResponse

# For reading .docx file and counting total number of words
from docx import Document
import re

# For counting word types (Considering Only English Texts)
from nltk.tag import pos_tag
from nltk.tokenize import word_tokenize



class MainView(TemplateView):
    template_name = 'docs/main.html'


def upload_document(request):
    if request.method == "POST":
        my_file = request.FILES.get('file')
        words_dict, total_words, tag_dict = count_words(my_file)
        return JsonResponse(
            {'total_words': f'Total Number of Words : {total_words}', 'words_dict': words_dict, 'tag_dict': tag_dict})
    return JsonResponse({'Upload': 'False'})


def count_words(file):
    words_dict = dict()
    tag_dict = {"ADJ": {"name": "Adjective", "count": 0},
                "ADP": {"name": "Adposition", "count": 0},
                "ADV": {"name": "Adverb", "count": 0},
                "AUX": {"name": "Auxiliary", "count": 0},
                "CCONJ": {"name": "Coordinating Conjunction", "count": 0},
                "DET": {"name": "Determiner", "count": 0},
                "NOUN": {"name": "Noun", "count": 0},
                "INTJ": {"name": "Interjection", "count": 0},
                "NUM": {"name": "Numeral", "count": 0},
                "PART": {"name": "Particle", "count": 0},
                "PRON": {"name": "Pronoun", "count": 0},
                "PROPN": {"name": "Proper Noun", "count": 0},
                "PUNCT": {"name": "Punctuation", "count": 0},
                "SCONJ": {"name": "Subordinating Conjunction", "count": 0},
                "SYM": {"name": "Symbol", "count": 0},
                "VERB": {"name": "Verb", "count": 0},
                "X": {"name": "Other", "count": 0},
                }
    total_words = 0
    document = Document(file)
    if document.tables:
        for table in document.tables:
            total_words += count_words_in_table(table)
    for para in document.paragraphs:
        string = para.text.lower()
        string = string.split(" ")

        for word in string:
            # removing some special characters
            word = re.sub("\ |\?|\,|\.|\!|\'|\/|\;|\:", '', word)

            if word == '':
                continue

            # Finding Type of Word (Considering Only English Texts)
            role = pos_tag(word_tokenize(f"{word}"), tagset='universal')
            # an example of role variable :  [('note', 'NOUN')] ---> A list which contains a tuple with two element
            try:
                # print(role)
                tag_dict[role[0][1]]['count'] += 1
            except KeyError:
                tag_dict['X']['count'] += 1
            except IndexError:
                continue
            except LookupError:
                print("Need to Download universal_tagset ")
                download_universal_tagset()
                print("universal_tagset has been downloaded, Continuing ")
            # Finding Type of Word has been finished
            try:
                words_dict[word] += 1
            except KeyError:
                words_dict[word] = 1
        total_words += len(para.text.split())

    if total_words >= 200:
        words_dict = dict(sorted(words_dict.items(), key=lambda kv: kv[1], reverse=True)[:20])
    else:
        words_dict = dict(sorted(words_dict.items(), key=lambda kv: kv[1], reverse=True)[:10])

    tag_dict = dict(sorted(tag_dict.items(), key=lambda kv: kv[1]['count'], reverse=True)[:10])

    return words_dict, total_words, tag_dict


def count_words_in_table(table):
    words = 0
    try:
        for row in table.rows:
            words += len(row.cells[0].text.split())
    except IndexError:
        print("Index Error")
    # print(words)
    return words


def download_universal_tagset():
    import nltk
    nltk.download('universal_tagset')
